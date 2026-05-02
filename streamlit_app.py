"""
RFP Compliance Analyzer — v6
==============================
Improvements over v5:
 (i)   Curated model allowlists per provider (3-4 known-good models each)
 (ii)  Live API verification: only show curated models the API key can access
 (iii) "Test API" button for one-click diagnosis of billing/rate/auth errors

Single-file Streamlit app. Session-based storage.
"""

import os
import hmac
import streamlit as st
import pandas as pd
import json
import re
import base64
import math
from io import BytesIO
from datetime import datetime

# ── optional imports (provider-specific) ──────────────────────────────────────
try:
    from anthropic import Anthropic
except ImportError:
    Anthropic = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    from groq import Groq
except ImportError:
    Groq = None

# ── document parsers ──────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from pypdf import PdfReader


# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════

st.set_page_config(
    page_title="RFP Compliance Analyzer",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Global CSS: guarantee scroll on long dropdowns
st.markdown(
    """
    <style>
    [role="listbox"],
    [data-baseweb="menu"] ul,
    [data-baseweb="popover"] ul {
        max-height: 400px !important;
        overflow-y: auto !important;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

BATCH_SIZE = 5


# ══════════════════════════════════════════════════════════════════════════════
# SECRETS & PASSWORD GATE
# ══════════════════════════════════════════════════════════════════════════════

def _get_secret(key: str):
    try:
        val = st.secrets.get(key)
        if val:
            return val
    except Exception:
        pass
    return os.getenv(key)


def check_password() -> bool:
    expected = _get_secret("APP_PASSWORD")
    if not expected:
        return True
    if st.session_state.get("password_correct", False):
        return True

    def _check():
        if hmac.compare_digest(
            st.session_state.get("password_input", ""),
            str(expected),
        ):
            st.session_state["password_correct"] = True
            try:
                del st.session_state["password_input"]
            except KeyError:
                pass
        else:
            st.session_state["password_correct"] = False

    st.markdown("### 🔒 RFP Compliance Analyzer")
    st.caption("Restricted access. Enter the password to continue.")
    st.text_input(
        "Password",
        type="password",
        on_change=_check,
        key="password_input",
    )
    if st.session_state.get("password_correct") is False:
        st.error("Incorrect password.")
    return False


# ══════════════════════════════════════════════════════════════════════════════
# CURATED MODEL ALLOWLISTS
# ══════════════════════════════════════════════════════════════════════════════
# Hand-picked, known-good models per provider. The dropdown shows the
# intersection of (allowlist) ∩ (live API list) — so the user only ever
# sees models that are both known-good AND verified accessible to their key.
# If the live fetch fails, the allowlist itself is shown.

ANTHROPIC_RECOMMENDED = [
    "claude-haiku-4-5-20251001",   # cheapest, fastest
    "claude-sonnet-4-6",            # balanced (default)
    "claude-opus-4-6",              # premium
    "claude-opus-4-7",              # latest premium
]

# All paid tier — picking only the latest stable family
OPENAI_RECOMMENDED = [
    "gpt-5.4-nano",   # cheapest
    "gpt-5.4-mini",   # cheap
    "gpt-5.4",        # balanced
    "gpt-5.5",        # premium
]

# Free-tier-friendly options first; pro is paid
GEMINI_RECOMMENDED = [
    "gemini-2.5-flash-lite",   # cheapest, free tier
    "gemini-2.5-flash",         # free tier (recommended)
    "gemini-2.5-pro",           # paid, best quality
]

# Only known-free Groq models — preview/paid models excluded
GROQ_RECOMMENDED = [
    "llama-3.1-8b-instant",          # fastest, free tier
    "llama-3.3-70b-versatile",       # balanced, free tier (recommended)
]


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_anthropic_models(api_key: str) -> list:
    if not api_key or Anthropic is None:
        return ANTHROPIC_RECOMMENDED
    try:
        client = Anthropic(api_key=api_key)
        page = client.models.list(limit=50)
        live_ids = {m.id for m in page.data}
        verified = [m for m in ANTHROPIC_RECOMMENDED if m in live_ids]
        return verified or ANTHROPIC_RECOMMENDED
    except Exception:
        return ANTHROPIC_RECOMMENDED


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_openai_models(api_key: str) -> list:
    if not api_key or OpenAI is None:
        return OPENAI_RECOMMENDED
    try:
        client = OpenAI(api_key=api_key)
        page = client.models.list()
        live_ids = {m.id for m in page.data}
        verified = [m for m in OPENAI_RECOMMENDED if m in live_ids]
        return verified or OPENAI_RECOMMENDED
    except Exception:
        return OPENAI_RECOMMENDED


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_gemini_models(api_key: str) -> list:
    if not api_key or genai is None:
        return GEMINI_RECOMMENDED
    try:
        genai.configure(api_key=api_key)
        live_ids = set()
        for m in genai.list_models():
            methods = getattr(m, "supported_generation_methods", []) or []
            if "generateContent" not in methods:
                continue
            name = (m.name or "").replace("models/", "")
            if name:
                live_ids.add(name)
        verified = [m for m in GEMINI_RECOMMENDED if m in live_ids]
        return verified or GEMINI_RECOMMENDED
    except Exception:
        return GEMINI_RECOMMENDED


@st.cache_data(ttl=3600, show_spinner=False)
def fetch_groq_models(api_key: str) -> list:
    if not api_key or Groq is None:
        return GROQ_RECOMMENDED
    try:
        client = Groq(api_key=api_key)
        page = client.models.list()
        live_ids = {m.id for m in page.data}
        verified = [m for m in GROQ_RECOMMENDED if m in live_ids]
        return verified or GROQ_RECOMMENDED
    except Exception:
        return GROQ_RECOMMENDED


def clear_model_caches():
    fetch_anthropic_models.clear()
    fetch_openai_models.clear()
    fetch_gemini_models.clear()
    fetch_groq_models.clear()


LLM_PROVIDERS = {
    "Anthropic Claude": {
        "fetch": fetch_anthropic_models,
        "secret_key": "ANTHROPIC_API_KEY",
        "native_pdf": True,
        "hint": "Best overall quality. Native PDF understanding. Paid only.",
    },
    "OpenAI": {
        "fetch": fetch_openai_models,
        "secret_key": "OPENAI_API_KEY",
        "native_pdf": False,
        "hint": "Strong reasoning. Paid only. PDFs are text-extracted first.",
    },
    "Google Gemini": {
        "fetch": fetch_gemini_models,
        "secret_key": "GOOGLE_API_KEY",
        "native_pdf": False,
        "hint": "Free tier on 2.5-flash and flash-lite. Large context window.",
    },
    "Groq (Open Source)": {
        "fetch": fetch_groq_models,
        "secret_key": "GROQ_API_KEY",
        "native_pdf": False,
        "hint": "Fully free tier. Fastest inference. Lower quality on legal nuance.",
    },
}


def get_models_for(provider_name: str) -> list:
    cfg = LLM_PROVIDERS[provider_name]
    api_key = _get_secret(cfg["secret_key"]) or ""
    return cfg["fetch"](api_key)


ANALYSIS_MODES = {
    "Contractual": """legal or commercial implications, including but not limited to:
- Standard terms and conditions / contract template provisions
- Liability, indemnification, warranty, and limitation clauses
- Intellectual property, data protection, and confidentiality provisions
- Payment terms, pricing, and fee arrangements
- Performance obligations, SLAs, and service level requirements
- Termination, renewal, and change-of-control clauses
- Governing law, jurisdiction, dispute resolution
- Scope of work items with contractual or legal implications
- Insurance, audit rights, compliance, and regulatory obligations""",

    "Technical Standards": """technical requirements, specifications, and standards, including:
- Functional and non-functional requirements
- Performance benchmarks, SLAs, and KPIs
- Integration requirements, APIs, interoperability
- Security standards, certifications, compliance frameworks
- Infrastructure, hosting, and deployment requirements
- Data formats, volumes, and processing requirements
- Support, maintenance, and response-time obligations
- Architecture, scalability, and availability targets""",

    "Custom": None,
}


# ══════════════════════════════════════════════════════════════════════════════
# FILE EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf_text(file_obj) -> str:
    file_obj.seek(0)
    reader = PdfReader(file_obj)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text and text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def extract_docx_text(file_obj) -> str:
    file_obj.seek(0)
    doc = DocxDocument(file_obj)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract_xlsx_text(file_obj) -> str:
    from openpyxl import load_workbook
    file_obj.seek(0)
    wb = load_workbook(file_obj, data_only=True)
    sections = []
    for name in wb.sheetnames:
        ws = wb[name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            if any(c is not None for c in row):
                rows.append(",".join(str(c) if c is not None else "" for c in row))
        if rows:
            sections.append(f"=== Sheet: {name} ===\n" + "\n".join(rows))
    return "\n\n".join(sections)


def prepare_file(uploaded_file, provider_name: str) -> dict:
    if uploaded_file is None:
        return None
    name = uploaded_file.name.lower()
    provider_cfg = LLM_PROVIDERS.get(provider_name, {})
    use_native_pdf = provider_cfg.get("native_pdf", False)
    uploaded_file.seek(0)

    if name.endswith(".pdf"):
        if use_native_pdf:
            b64 = base64.standard_b64encode(uploaded_file.read()).decode("utf-8")
            return {
                "type": "document",
                "payload": {
                    "type": "document",
                    "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
                },
            }
        else:
            return {"type": "text", "payload": extract_pdf_text(uploaded_file)}

    if name.endswith(".docx"):
        return {"type": "text", "payload": extract_docx_text(uploaded_file)}

    if name.endswith((".xlsx", ".xls", ".xlsm")):
        return {"type": "text", "payload": extract_xlsx_text(uploaded_file)}

    uploaded_file.seek(0)
    return {"type": "text", "payload": uploaded_file.read().decode("utf-8", errors="ignore")}


# ══════════════════════════════════════════════════════════════════════════════
# LLM ABSTRACTION LAYER
# ══════════════════════════════════════════════════════════════════════════════

def call_llm(system: str, user_blocks, max_tokens: int = 4000) -> str:
    provider = st.session_state.get("llm_provider", "Anthropic Claude")
    model = st.session_state.get("llm_model", "claude-sonnet-4-6")

    if provider == "Anthropic Claude":
        key = _get_secret("ANTHROPIC_API_KEY")
        if not key:
            st.error("ANTHROPIC_API_KEY missing from Streamlit secrets.")
            st.stop()
        client = Anthropic(api_key=key)
        resp = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": user_blocks}],
        )
        return "".join(b.text for b in resp.content if hasattr(b, "text"))

    if provider == "OpenAI":
        key = _get_secret("OPENAI_API_KEY")
        if not key:
            st.error("OPENAI_API_KEY missing from Streamlit secrets.")
            st.stop()
        client = OpenAI(api_key=key)
        text = user_blocks if isinstance(user_blocks, str) else _blocks_to_text(user_blocks)
        resp = client.chat.completions.create(
            model=model,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": text},
            ],
        )
        return resp.choices[0].message.content or ""

    if provider == "Google Gemini":
        key = _get_secret("GOOGLE_API_KEY")
        if not key:
            st.error("GOOGLE_API_KEY missing from Streamlit secrets.")
            st.stop()
        genai.configure(api_key=key)
        text = user_blocks if isinstance(user_blocks, str) else _blocks_to_text(user_blocks)
        m = genai.GenerativeModel(model_name=model, system_instruction=system)
        resp = m.generate_content(
            text,
            generation_config=genai.GenerationConfig(max_output_tokens=max_tokens),
        )
        return resp.text or ""

    if provider == "Groq (Open Source)":
        key = _get_secret("GROQ_API_KEY")
        if not key:
            st.error("GROQ_API_KEY missing from Streamlit secrets.")
            st.stop()
        client = Groq(api_key=key)
        text = user_blocks if isinstance(user_blocks, str) else _blocks_to_text(user_blocks)
        resp = client.chat.completions.create(
            model=model,
            max_tokens=max_tokens,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": text},
            ],
        )
        return resp.choices[0].message.content or ""

    raise ValueError(f"Unknown provider: {provider}")


def _blocks_to_text(blocks) -> str:
    if isinstance(blocks, str):
        return blocks
    parts = []
    for b in blocks:
        if isinstance(b, dict):
            if b.get("type") == "text":
                parts.append(b["text"])
            elif b.get("type") == "document":
                parts.append("[PDF document — text extraction used for this provider]")
        else:
            parts.append(str(b))
    return "\n\n".join(parts)


def diagnose_error(err: Exception) -> tuple:
    """
    Inspect an exception from call_llm and return (icon, category, hint).
    Used by the Test API button to give actionable diagnosis.
    """
    msg = str(err).lower()
    if "credit" in msg or "balance" in msg or "insufficient_quota" in msg:
        return ("💳", "Billing",
                "Your account has zero credits. Top up at the provider's billing page.")
    if "rate" in msg and "limit" in msg:
        return ("⏰", "Rate limit",
                "Too many requests in a short window. Wait 60s and retry, or pick a smaller model.")
    if "quota" in msg or "limit: 0" in msg:
        return ("📊", "Quota",
                "Your account/region has no quota allocated for this model. "
                "Check provider console; may need billing enabled.")
    if "401" in msg or "unauthor" in msg or "invalid" in msg and "key" in msg:
        return ("🔑", "Auth",
                "API key is invalid, expired, or revoked. "
                "Generate a new key in the provider console and update Streamlit secrets.")
    if "404" in msg or "not_found" in msg or "does not exist" in msg:
        return ("❌", "Model not accessible",
                "The selected model is not available to your API tier. "
                "Try a different model from the dropdown.")
    if "tokens" in msg and ("limit" in msg or "exceed" in msg or "exhausted" in msg):
        return ("📊", "Token limit",
                "Daily/monthly token cap reached. Try again tomorrow, "
                "switch to a smaller model, or upgrade tier.")
    if "context" in msg and ("length" in msg or "window" in msg):
        return ("📏", "Context too long",
                "Document is larger than this model's context window. "
                "Try a model with bigger context (Gemini 2.5 Pro, Claude Opus).")
    return ("✗", "Unknown error", str(err)[:300])


def extract_json_array(text: str):
    if not text:
        return None
    clean = re.sub(r"```json|```", "", text).strip()
    start = clean.find("[")
    if start == -1:
        return None

    end = clean.rfind("]")
    if end > start:
        try:
            return json.loads(clean[start : end + 1])
        except json.JSONDecodeError:
            pass

    body = clean[start + 1:]
    depth = 0
    in_str = False
    esc = False
    last_good = -1
    for i, ch in enumerate(body):
        if esc:
            esc = False
            continue
        if ch == "\\":
            esc = True
            continue
        if ch == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                last_good = i

    if last_good == -1:
        return None
    try:
        return json.loads("[" + body[: last_good + 1] + "]")
    except json.JSONDecodeError:
        return None


# ══════════════════════════════════════════════════════════════════════════════
# CORRECTIONS (TEACHING LOOP)
# ══════════════════════════════════════════════════════════════════════════════

def build_corrections_prompt(corrections: list, mode: str, limit: int = 8) -> str:
    relevant = [c for c in corrections if c.get("mode") == mode]
    if not relevant:
        relevant = corrections
    recent = relevant[-limit:]
    if not recent:
        return ""

    lines = [
        "═══ LEARNING FROM PAST CORRECTIONS ═══",
        "The following examples show cases where an earlier AI analysis was",
        "overridden by the user. Apply these standards consistently.\n",
    ]
    for i, c in enumerate(recent, 1):
        lines.append(f"── Correction {i} ──")
        lines.append(f"RFP clause   : \"{c.get('clauseText', '')[:250]}\"")
        if c.get("playbookClause"):
            lines.append(f"Playbook ref : \"{c['playbookClause'][:250]}\"")
        orig_cls = c.get("original_classification", "?")
        corr_cls = c.get("corrected_classification", "?")
        if orig_cls != corr_cls:
            lines.append(f"AI classified : {orig_cls}  →  CORRECTED TO: {corr_cls}")
        else:
            lines.append(f"Classification: {corr_cls} (confirmed correct)")
        lines.append(f"Correct reason: {c.get('corrected_reason', '')}")
        if c.get("corrected_alternative"):
            lines.append(f"Correct alt.  : {c.get('corrected_alternative', '')}")
        lines.append("")
    lines.append("═══ END OF CORRECTIONS ═══\n")
    return "\n".join(lines)


def capture_corrections_from_edits(original: list, edited: list, mode: str) -> list:
    orig_map = {str(r.get("id", i)): r for i, r in enumerate(original)}
    corrections = []
    for i, row in enumerate(edited):
        rid = str(row.get("id", i))
        orig = orig_map.get(rid)
        if not orig:
            continue
        changed = (
            orig.get("classification") != row.get("classification")
            or orig.get("reason") != row.get("reason")
            or orig.get("alternative") != row.get("alternative")
        )
        if changed:
            corrections.append(
                {
                    "clauseText": orig.get("clauseText", ""),
                    "section": orig.get("section", ""),
                    "playbookClause": orig.get("playbookClause", ""),
                    "original_classification": orig.get("classification", ""),
                    "original_reason": orig.get("reason", ""),
                    "original_alternative": orig.get("alternative", ""),
                    "corrected_classification": row.get("classification", ""),
                    "corrected_reason": row.get("reason", ""),
                    "corrected_alternative": row.get("alternative", ""),
                    "mode": mode,
                    "timestamp": datetime.now().isoformat(),
                }
            )
    return corrections


# ══════════════════════════════════════════════════════════════════════════════
# ANALYSIS PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

def extract_clauses(rfp_file, provider: str, mode: str, custom_instructions: str) -> list:
    rfp_data = prepare_file(rfp_file, provider)
    focus = custom_instructions if mode == "Custom" else ANALYSIS_MODES[mode]
    if not focus:
        focus = "clauses relevant to the user's comparison objective"

    system = (
        "You are a senior commercial and legal analyst extracting clauses from RFP "
        "documents for vendor review. Return ONLY valid JSON arrays. No preamble, "
        "no commentary, no markdown fences."
    )

    instruction = f"""Analyze the RFP document. Extract every distinct clause, term, provision, or
requirement that carries {focus}

For each item return:
 - "id": sequential string ("1", "2", …)
 - "section": nearest heading or clause number from the document
 - "clauseText": verbatim quote (max 400 chars; summarise only if essential)

Extract 10–30 most material items. Skip neutral boilerplate (cover page, index, contacts).

Return ONLY a JSON array:
[{{"id":"1","section":"Clause 3.1","clauseText":"..."}}, ...]"""

    if rfp_data["type"] == "document":
        user_blocks = [rfp_data["payload"], {"type": "text", "text": instruction}]
    else:
        user_blocks = f"RFP DOCUMENT:\n\n{rfp_data['payload']}\n\n---\n\n{instruction}"

    raw = call_llm(system, user_blocks, max_tokens=8000)
    clauses = extract_json_array(raw)
    if not clauses:
        preview = (raw or "")[:500]
        raise ValueError(
            f"Could not extract clauses.\n\nModel response preview:\n{preview}\n\n"
            "Common causes: scanned PDF without OCR, empty document, or unexpected format."
        )
    return clauses


def analyze_batch(
    clauses: list,
    playbook_file,
    provider: str,
    mode: str,
    custom_instructions: str,
    corrections: list,
) -> list:
    pb_data = prepare_file(playbook_file, provider)
    focus = custom_instructions if mode == "Custom" else mode.lower()
    corrections_block = build_corrections_prompt(corrections, mode)

    system = (
        "You are a senior commercial and legal analyst classifying RFP clauses against "
        "a vendor's playbook/reference standard.\n\n"
        "Classification rules:\n"
        "  C  (Comply)         — clause aligns with playbook; acceptable as-is.\n"
        "  NC (Non-Compliant)  — clause is detrimental or violates playbook; cannot "
        "be accepted without substantial rewrite.\n"
        "  PC (Partially Compliant) — some elements acceptable, others need modification. "
        "Use only when there is a genuine split — not as a default middle ground.\n\n"
        "For NC and PC: give a specific reason and a constructive alternative.\n"
        "For C: reason can be brief; alternative must be empty string.\n\n"
        "IMPORTANT — also identify the specific playbook clause or provision that "
        "informed your classification. Quote or summarise it concisely.\n\n"
        "Return ONLY a valid JSON array, no preamble, no markdown fences."
    )

    extra = f"ADDITIONAL INSTRUCTIONS: {custom_instructions}\n\n" if custom_instructions else ""

    instruction = f"""{corrections_block}{extra}Analyze the {len(clauses)} RFP clauses below against the playbook/reference standard.

Context: {focus}

Clauses:
{json.dumps(clauses, indent=2)}

Return ONE object per clause (matching IDs):
[{{
  "id": "1",
  "playbookClause": "Section X of playbook: ...",
  "classification": "C|PC|NC",
  "reason": "...",
  "alternative": ""
}}, ...]"""

    if pb_data["type"] == "document":
        user_blocks = [pb_data["payload"], {"type": "text", "text": instruction}]
    else:
        user_blocks = f"PLAYBOOK / REFERENCE STANDARD:\n\n{pb_data['payload']}\n\n---\n\n{instruction}"

    raw = call_llm(system, user_blocks, max_tokens=4000)
    analyses = extract_json_array(raw)
    if not analyses:
        return [
            {"id": c["id"], "playbookClause": "", "classification": "",
             "reason": "(AI analysis failed — review manually)", "alternative": ""}
            for c in clauses
        ]
    return analyses


def run_full_analysis(
    rfp_file, playbook_file, provider, mode, custom_instructions, corrections, progress_cb
) -> list:
    progress_cb(5, "Extracting clauses from RFP…")
    clauses = extract_clauses(rfp_file, provider, mode, custom_instructions)
    progress_cb(20, f"Found {len(clauses)} clauses. Analyzing against playbook…")

    results = []
    for i in range(0, len(clauses), BATCH_SIZE):
        batch = clauses[i : i + BATCH_SIZE]
        playbook_file.seek(0)
        analyses = analyze_batch(
            batch, playbook_file, provider, mode, custom_instructions, corrections
        )
        for clause in batch:
            analysis = next(
                (a for a in analyses if str(a.get("id")) == str(clause["id"])), {}
            )
            results.append(
                {
                    "id": clause["id"],
                    "section": clause.get("section", ""),
                    "clauseText": clause.get("clauseText", ""),
                    "playbookClause": analysis.get("playbookClause", ""),
                    "classification": analysis.get("classification", ""),
                    "reason": analysis.get("reason", ""),
                    "alternative": analysis.get("alternative", ""),
                }
            )
        processed = min(i + BATCH_SIZE, len(clauses))
        pct = 20 + int((processed / len(clauses)) * 75)
        progress_cb(pct, f"Analyzed {processed} of {len(clauses)} clauses…")

    progress_cb(100, "Analysis complete.")
    return results


# ══════════════════════════════════════════════════════════════════════════════
# EXPORTS
# ══════════════════════════════════════════════════════════════════════════════

def _counts(results):
    c = {"C": 0, "PC": 0, "NC": 0, "Pending": 0}
    for r in results:
        cls = r.get("classification", "")
        if cls == "C":
            c["C"] += 1
        elif cls == "PC":
            c["PC"] += 1
        elif cls == "NC":
            c["NC"] += 1
        else:
            c["Pending"] += 1
    return c


STATUS_LABELS = {
    "C": "C (Comply)",
    "NC": "NC (Non-Compliant)",
    "PC": "PC (Partially Compliant)",
}

COLUMN_ORDER = ["#", "Section / Clause", "RFP Clause", "Playbook Reference",
                "Status", "Comments", "Suggested Alternative"]


def _safe_str(value) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and math.isnan(value):
        return ""
    return str(value)


def export_to_word(results, metadata, title, subtitle="", prepared_by="") -> BytesIO:
    doc = DocxDocument()
    doc.add_heading(_safe_str(title), level=0)
    if subtitle:
        p = doc.add_paragraph(_safe_str(subtitle))
        if p.runs:
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    mp = doc.add_paragraph()
    for lbl, val in [
        ("RFP: ", metadata.get("rfp_name", "—")),
        ("\nPlaybook: ", metadata.get("pb_name", "—")),
        ("\nMode: ", metadata.get("mode", "—")),
        ("\nProvider: ", f"{metadata.get('provider','')} / {metadata.get('model','')}"),
        ("\nDate: ", datetime.now().strftime("%d %B %Y")),
    ]:
        r = mp.add_run(_safe_str(lbl))
        r.bold = True
        r.font.size = Pt(10)
        r = mp.add_run(_safe_str(val))
        r.font.size = Pt(10)

    if prepared_by:
        r = mp.add_run("\nPrepared by: ")
        r.bold = True
        r.font.size = Pt(10)
        mp.add_run(_safe_str(prepared_by)).font.size = Pt(10)

    counts = _counts(results)
    total = len(results)

    doc.add_heading("Summary", level=1)
    st_table = doc.add_table(rows=1, cols=3)
    st_table.style = "Light Grid Accent 1"
    for cell, h in zip(st_table.rows[0].cells, ["Classification", "Count", "%"]):
        cell.text = h

    for lbl, key in [
        ("Comply (C)", "C"),
        ("Partially Compliant (PC)", "PC"),
        ("Non-Compliant (NC)", "NC"),
    ]:
        row = st_table.add_row().cells
        row[0].text = lbl
        row[1].text = str(counts[key])
        row[2].text = f"{round(counts[key]/total*100) if total else 0}%"

    if counts["Pending"]:
        row = st_table.add_row().cells
        row[0].text = "Pending Review"
        row[1].text = str(counts["Pending"])
        row[2].text = f"{round(counts['Pending']/total*100) if total else 0}%"

    tot_row = st_table.add_row().cells
    tot_row[0].text = "Total"
    tot_row[1].text = str(total)
    tot_row[2].text = "100%"

    doc.add_heading("Clause Analysis", level=1)
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    table.autofit = False

    col_widths = [Inches(0.32), Inches(1.1), Inches(1.9), Inches(1.5),
                  Inches(0.82), Inches(1.3), Inches(1.3)]

    for cell, h, w in zip(table.rows[0].cells, COLUMN_ORDER, col_widths):
        cell.text = h
        cell.width = w
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)

    status_colors_rgb = {
        "C": RGBColor(0x2D, 0x50, 0x16),
        "NC": RGBColor(0x8B, 0x2A, 0x1F),
        "PC": RGBColor(0xA0, 0x64, 0x0C),
    }

    for i, r in enumerate(results, 1):
        cls = r.get("classification", "") or ""
        row = table.add_row().cells
        vals = [
            str(i),
            _safe_str(r.get("section")),
            _safe_str(r.get("clauseText")),
            _safe_str(r.get("playbookClause")),
            STATUS_LABELS.get(cls, "Pending"),
            _safe_str(r.get("reason")),
            _safe_str(r.get("alternative")),
        ]
        for cell, val, w in zip(row, vals, col_widths):
            cell.text = val
            cell.width = w
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
        if cls in status_colors_rgb:
            for p in row[4].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = status_colors_rgb[cls]

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def export_to_excel(results, metadata, title, subtitle="", prepared_by="") -> BytesIO:
    wb = Workbook()
    counts = _counts(results)
    total = len(results)

    ws = wb.active
    ws.title = "Summary"
    ws["A1"] = _safe_str(title)
    ws["A1"].font = Font(name="Georgia", size=16, bold=True)
    if subtitle:
        ws["A2"] = _safe_str(subtitle)
        ws["A2"].font = Font(name="Calibri", size=11, color="7A7570")

    label_font = Font(name="Calibri", size=10, bold=True)
    for r, (lbl, val) in enumerate(
        [
            ("RFP:", metadata.get("rfp_name", "")),
            ("Playbook:", metadata.get("pb_name", "")),
            ("Mode:", metadata.get("mode", "")),
            ("Provider:", metadata.get("provider", "")),
            ("Model:", metadata.get("model", "")),
            ("Date:", datetime.now().strftime("%Y-%m-%d")),
            ("Prepared by:", prepared_by or "—"),
        ],
        start=4,
    ):
        ws.cell(row=r, column=1, value=lbl).font = label_font
        ws.cell(row=r, column=2, value=_safe_str(val))

    row = 12
    for ci, lbl in enumerate(["Classification", "Count", "%"], start=1):
        ws.cell(row=row, column=ci, value=lbl).font = Font(bold=True)
    row += 1

    for lbl, key, color in [
        ("Comply (C)", "C", "2D5016"),
        ("Partially Compliant (PC)", "PC", "A0640C"),
        ("Non-Compliant (NC)", "NC", "8B2A1F"),
        ("Pending Review", "Pending", "7A7570"),
    ]:
        if counts[key] == 0 and key == "Pending":
            continue
        ws.cell(row=row, column=1, value=lbl).font = Font(color=color, bold=True)
        ws.cell(row=row, column=2, value=counts[key])
        ws.cell(row=row, column=3,
                value=f"{round(counts[key]/total*100) if total else 0}%")
        row += 1

    ws.cell(row=row, column=1, value="Total").font = Font(bold=True)
    ws.cell(row=row, column=2, value=total).font = Font(bold=True)
    ws.cell(row=row, column=3, value="100%").font = Font(bold=True)

    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 12

    wa = wb.create_sheet("Analysis")
    headers = COLUMN_ORDER
    widths = [5, 24, 45, 35, 18, 38, 38]
    hdr_fill = PatternFill(start_color="2D5016", end_color="2D5016", fill_type="solid")
    hdr_font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )
    for ci, (h, w) in enumerate(zip(headers, widths), 1):
        cell = wa.cell(row=1, column=ci, value=h)
        cell.fill = hdr_fill
        cell.font = hdr_font
        cell.border = border
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        wa.column_dimensions[get_column_letter(ci)].width = w

    row_fills = {
        "C": PatternFill(start_color="F4F7EF", end_color="F4F7EF", fill_type="solid"),
        "NC": PatternFill(start_color="FAECEB", end_color="FAECEB", fill_type="solid"),
        "PC": PatternFill(start_color="FCF5E6", end_color="FCF5E6", fill_type="solid"),
    }
    status_colors = {"C": "2D5016", "NC": "8B2A1F", "PC": "A0640C"}

    for ri, r in enumerate(results, 2):
        cls = r.get("classification", "")
        values = [
            ri - 1,
            _safe_str(r.get("section")),
            _safe_str(r.get("clauseText")),
            _safe_str(r.get("playbookClause")),
            STATUS_LABELS.get(cls, "Pending"),
            _safe_str(r.get("reason")),
            _safe_str(r.get("alternative")),
        ]
        fill = row_fills.get(cls)
        for ci, val in enumerate(values, 1):
            cell = wa.cell(row=ri, column=ci, value=val)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if fill:
                cell.fill = fill
        if cls in status_colors:
            wa.cell(row=ri, column=5).font = Font(bold=True, color=status_colors[cls])

    wa.freeze_panes = "A2"

    bio = BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio


# ══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ══════════════════════════════════════════════════════════════════════════════

def init_state():
    defaults = {
        "llm_provider": "Anthropic Claude",
        "llm_model": "claude-sonnet-4-6",
        "results": None,
        "original_results": None,
        "metadata": None,
        "analysis_title": "",
        "analysis_mode": "Contractual",
        "corrections": [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


# ══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════

def render_sidebar():
    with st.sidebar:
        st.markdown("### 📋 RFP Compliance Analyzer")
        st.caption("v6 · Session-based")
        st.divider()

        # ── LLM Provider ──────────────────────────────────────────────────────
        st.markdown("#### 🤖 AI Provider")
        provider = st.selectbox(
            "Provider",
            options=list(LLM_PROVIDERS.keys()),
            index=list(LLM_PROVIDERS.keys()).index(st.session_state["llm_provider"]),
            label_visibility="collapsed",
        )
        if provider != st.session_state["llm_provider"]:
            st.session_state["llm_provider"] = provider
            st.session_state["llm_model"] = None

        cfg = LLM_PROVIDERS[provider]

        with st.spinner("Loading models..."):
            models = get_models_for(provider)

        current = st.session_state.get("llm_model")
        if current not in models:
            current = models[0] if models else None
            st.session_state["llm_model"] = current

        col_a, col_b = st.columns([4, 1])
        with col_a:
            model = st.selectbox(
                "Model",
                options=models,
                index=models.index(current) if current in models else 0,
                label_visibility="collapsed",
            )
        with col_b:
            if st.button("🔄", help="Refresh model list from provider", use_container_width=True):
                clear_model_caches()
                st.rerun()

        st.session_state["llm_model"] = model
        st.caption(cfg["hint"])

        key_present = bool(_get_secret(cfg["secret_key"]))
        if key_present:
            st.success(f"✓ {cfg['secret_key']} found · {len(models)} model(s) verified")
        else:
            st.error(f"✗ {cfg['secret_key']} missing — using curated default list")

        # ── Test API button ──────────────────────────────────────────────────
        if st.button("🧪 Test API", help="Send a tiny test request to verify the key + model work",
                     use_container_width=True):
            if not key_present:
                st.error("Cannot test — API key not configured.")
            else:
                with st.spinner(f"Testing {model}..."):
                    try:
                        response = call_llm(
                            system="Reply with only the word: ok",
                            user_blocks="Say ok",
                            max_tokens=10,
                        )
                        if response and response.strip():
                            st.success(f"✓ Working. Got: '{response.strip()[:40]}'")
                        else:
                            st.warning("⚠ Empty response (model may need different params)")
                    except Exception as e:
                        icon, category, hint = diagnose_error(e)
                        st.error(f"{icon} **{category}**\n\n{hint}")
                        with st.expander("Full error"):
                            st.code(str(e)[:1000])

        if not cfg["native_pdf"]:
            st.info("ℹ️ PDFs will be text-extracted (this provider does not support native PDF).")

        st.divider()

        # ── Teaching Corrections ──────────────────────────────────────────────
        corrections = st.session_state["corrections"]
        st.markdown(f"#### 🎓 Teaching Examples ({len(corrections)})")
        if corrections:
            st.caption(
                f"{len(corrections)} correction(s) saved. These are injected into "
                "every new analysis to calibrate the AI to your standards."
            )
            corr_json = json.dumps(corrections, indent=2)
            st.download_button(
                "💾 Download corrections (.json)",
                data=corr_json,
                file_name=f"rfp_corrections_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                use_container_width=True,
            )
            if st.button("🗑️ Clear all corrections", use_container_width=True):
                st.session_state["corrections"] = []
                st.rerun()
            with st.expander("View corrections"):
                for i, c in enumerate(corrections, 1):
                    st.markdown(
                        f"**{i}.** `{c.get('original_classification','?')}` → "
                        f"`{c.get('corrected_classification','?')}` — "
                        f"*{c.get('section', 'Unknown section')}*"
                    )
        else:
            st.caption(
                "No corrections yet. Run an analysis, edit the results, then "
                "click **Capture corrections** to teach the AI your standards."
            )

        uploaded_corr = st.file_uploader(
            "Upload corrections from previous session",
            type=["json"],
            key="corrections_upload",
        )
        if uploaded_corr:
            try:
                loaded = json.loads(uploaded_corr.read().decode("utf-8"))
                if isinstance(loaded, list):
                    st.session_state["corrections"] = loaded
                    st.success(f"Loaded {len(loaded)} correction(s).")
                    st.rerun()
            except Exception as e:
                st.error(f"Could not load corrections: {e}")

        st.divider()
        st.markdown("#### How it works")
        st.markdown(
            "1. Upload **RFP** and **Playbook**\n"
            "2. Pick **mode** and **provider**\n"
            "3. **Run Analysis** — AI extracts and classifies\n"
            "4. **Edit** any wrong cells in the table\n"
            "5. **Capture corrections** → AI learns your standards\n"
            "6. **Download** as Word or Excel"
        )

        st.divider()
        st.caption(
            "**Privacy:** Documents are sent to the selected provider's API. "
            "No data is stored server-side beyond the current session."
        )

        st.divider()
        if st.button("🔄 Clear session", use_container_width=True):
            for k in [k for k in st.session_state if k != "corrections"]:
                del st.session_state[k]
            st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# NEW ANALYSIS VIEW
# ══════════════════════════════════════════════════════════════════════════════

def render_new_analysis():
    st.markdown("### Step 1 · Upload documents")
    col1, col2 = st.columns(2)
    with col1:
        rfp = st.file_uploader(
            "**RFP document**",
            type=["pdf", "docx", "xlsx", "xls", "xlsm", "txt", "md"],
            key="rfp_upload",
        )
    with col2:
        pb = st.file_uploader(
            "**Playbook / reference standard**",
            type=["pdf", "docx", "xlsx", "xls", "xlsm", "txt", "md"],
            key="pb_upload",
        )

    st.markdown("### Step 2 · Analysis mode")
    mode = st.radio(
        "Mode",
        options=list(ANALYSIS_MODES.keys()),
        horizontal=True,
        label_visibility="collapsed",
        index=list(ANALYSIS_MODES.keys()).index(
            st.session_state.get("analysis_mode", "Contractual")
        ),
    )
    st.session_state["analysis_mode"] = mode

    hints = {
        "Contractual": "T&Cs, liability, IP, data, payment, termination — legal and commercial clauses.",
        "Technical Standards": "Specifications, benchmarks, integrations, performance metrics.",
        "Custom": "Write your own comparison instructions below.",
    }
    st.caption(hints[mode])

    if mode == "Custom":
        custom_instructions = st.text_area(
            "Comparison instructions",
            placeholder="Describe what to compare and how to classify clauses...",
            height=120,
            key="custom_instructions_input",
        )
    else:
        custom_instructions = st.text_area(
            "Additional instructions (optional)",
            placeholder="Any extra guidance for the AI (optional)...",
            height=80,
            key="custom_instructions_input",
        )

    st.markdown("### Step 3 · Run")
    if not (rfp and pb):
        st.info("Upload both an RFP and a playbook to proceed.")

    if st.button(
        "🚀 Run Analysis",
        type="primary",
        disabled=not (rfp and pb),
    ):
        bar = st.progress(0)
        status = st.empty()

        def update(pct, msg):
            bar.progress(pct / 100)
            status.markdown(f"*{msg}*")

        try:
            results = run_full_analysis(
                rfp_file=rfp,
                playbook_file=pb,
                provider=st.session_state["llm_provider"],
                mode=mode,
                custom_instructions=custom_instructions,
                corrections=st.session_state["corrections"],
                progress_cb=update,
            )
            st.session_state["results"] = results
            st.session_state["original_results"] = [r.copy() for r in results]
            st.session_state["metadata"] = {
                "rfp_name": rfp.name,
                "pb_name": pb.name,
                "mode": mode,
                "provider": st.session_state["llm_provider"],
                "model": st.session_state["llm_model"],
                "date": datetime.now().isoformat(),
            }
            if not st.session_state["analysis_title"]:
                st.session_state["analysis_title"] = (
                    rfp.name.rsplit(".", 1)[0] + " — Compliance Analysis"
                )
            bar.empty()
            status.empty()
            st.success(f"✓ Done — {len(results)} clauses extracted and classified.")
            st.rerun()
        except Exception as e:
            bar.empty()
            status.empty()
            icon, category, hint = diagnose_error(e)
            st.error(f"**Analysis failed — {icon} {category}**\n\n{hint}")
            with st.expander("Full error trace"):
                st.exception(e)


# ══════════════════════════════════════════════════════════════════════════════
# RESULTS VIEW
# ══════════════════════════════════════════════════════════════════════════════

def render_results():
    results = st.session_state["results"]
    original = st.session_state.get("original_results", [])
    meta = st.session_state["metadata"]
    mode = meta.get("mode", "Contractual")

    st.session_state["analysis_title"] = st.text_input(
        "Analysis title",
        value=st.session_state["analysis_title"],
    )
    st.caption(
        f"**RFP:** {meta['rfp_name']} · **Playbook:** {meta['pb_name']} · "
        f"**Mode:** {meta['mode']} · **Provider:** {meta['provider']} / {meta['model']}"
    )

    counts = _counts(results)
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Total", len(results))
    c2.metric("✅ C", counts["C"])
    c3.metric("🟡 PC", counts["PC"])
    c4.metric("🔴 NC", counts["NC"])
    c5.metric("⏳ Pending", counts["Pending"])

    st.markdown("#### Clause matrix")
    st.caption(
        "Edit any cell directly. The **Playbook Reference** column shows which playbook "
        "provision the AI used for each classification. Change a classification via its "
        "dropdown. When done editing, click **Capture corrections** to teach the AI."
    )

    df = pd.DataFrame(results)

    edited_df = st.data_editor(
        df,
        column_config={
            "id": st.column_config.TextColumn("#", width="small", disabled=True),
            "section": st.column_config.TextColumn("Section / Clause", width="medium"),
            "clauseText": st.column_config.TextColumn("RFP Clause", width="large"),
            "playbookClause": st.column_config.TextColumn(
                "Playbook Reference", width="large",
                help="The specific playbook/standard provision the AI compared against."
            ),
            "classification": st.column_config.SelectboxColumn(
                "Status",
                options=["C", "PC", "NC", ""],
                required=False,
                width="small",
            ),
            "reason": st.column_config.TextColumn("Reason / Comments", width="large"),
            "alternative": st.column_config.TextColumn("Suggested Alternative", width="large"),
        },
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        key="results_editor",
    )

    edited_records = edited_df.to_dict("records")
    st.session_state["results"] = edited_records

    new_corrections = capture_corrections_from_edits(original, edited_records, mode)
    if new_corrections:
        st.markdown(
            f"💡 **{len(new_corrections)} edit(s) detected** that differ from the AI's "
            "original output."
        )
        if st.button(
            f"🎓 Capture {len(new_corrections)} correction(s) as teaching examples",
            type="primary",
        ):
            st.session_state["corrections"].extend(new_corrections)
            st.session_state["original_results"] = [r.copy() for r in edited_records]
            st.success(
                f"✓ Saved {len(new_corrections)} correction(s). "
                f"Total: {len(st.session_state['corrections'])}. "
                "These will be applied to your next analysis."
            )
            st.rerun()

    st.divider()
    st.markdown("#### Export")
    col1, col2 = st.columns(2)
    with col1:
        subtitle = st.text_input("Subtitle (optional)", placeholder="Vendor Response Draft v2")
    with col2:
        prepared_by = st.text_input("Prepared by", placeholder="Name or team")

    export_title = st.session_state["analysis_title"] or "RFP Compliance Analysis"
    safe_name = re.sub(r"[^a-zA-Z0-9]+", "_", export_title).strip("_")
    date_stamp = datetime.now().strftime("%Y%m%d")

    dl1, dl2 = st.columns(2)
    with dl1:
        try:
            word_bio = export_to_word(edited_records, meta, export_title, subtitle, prepared_by)
            st.download_button(
                "📝 Download Word",
                data=word_bio,
                file_name=f"{safe_name}_{date_stamp}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"Could not build Word document: {e}")

    with dl2:
        try:
            excel_bio = export_to_excel(edited_records, meta, export_title, subtitle, prepared_by)
            st.download_button(
                "📊 Download Excel",
                data=excel_bio,
                file_name=f"{safe_name}_{date_stamp}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"Could not build Excel document: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    init_state()

    if not check_password():
        st.stop()

    render_sidebar()

    st.title("RFP Compliance Analyzer")
    st.markdown(
        "*Clause-level C / NC / PC analysis — with playbook references, multi-LLM "
        "support, and an AI teaching loop.*"
    )
    st.divider()

    if st.session_state["results"] is None:
        render_new_analysis()
    else:
        render_results()
        st.divider()
        if st.button("🔄 Start a new analysis"):
            st.session_state["results"] = None
            st.session_state["original_results"] = None
            st.session_state["metadata"] = None
            st.session_state["analysis_title"] = ""
            st.rerun()


if __name__ == "__main__":
    main()
